"""
Gerador de documentação do Sistema de Controle de Estoque v3.1.0
Salva como .docx na área de trabalho.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

DESKTOP = r"C:\Users\wemerson - T.I\Desktop"
OUTPUT  = os.path.join(DESKTOP, "Documentacao_SistemaEstoque_v3.1.docx")

# ── Cores ──────────────────────────────────────────────────────────────────────
VERDE_ESCURO  = RGBColor(0x1E, 0x3A, 0x34)
VERDE_MEDIO   = RGBColor(0x2E, 0x7D, 0x32)
OURO          = RGBColor(0xB8, 0x95, 0x2A)
CINZA_ESCURO  = RGBColor(0x37, 0x47, 0x4F)
CINZA_CLARO   = RGBColor(0xF5, 0xF5, 0xF5)
BRANCO        = RGBColor(0xFF, 0xFF, 0xFF)
VERMELHO      = RGBColor(0xC6, 0x28, 0x28)

doc = Document()

# ── Configuração de página ─────────────────────────────────────────────────────
for sec in doc.sections:
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin  = sec.bottom_margin = Cm(2)


# ── Helpers ────────────────────────────────────────────────────────────────────
def shading(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def cell_border(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "CCCCCC")
        tcBorders.append(b)
    tcPr.append(tcBorders)

def heading(text, level=1, color=VERDE_ESCURO, space_before=16, space_after=6):
    p    = doc.add_paragraph()
    run  = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt({1:20, 2:15, 3:13, 4:11}.get(level, 11))
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if level == 1:
        pf.keep_with_next = True
    return p

def body(text, bold=False, italic=False, color=None, size=10, space_after=4):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p

def bullet(text, level=0, color=None):
    p   = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.6)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    return p

def code_block(text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = CINZA_ESCURO
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    return p

def divider():
    p  = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = pf.space_after = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "C5C5C5")
    pBdr.append(bot)
    pPr.append(pBdr)

def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # cabeçalho
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        c   = hdr.cells[i]
        shading(c, "1E3A34")
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = BRANCO
        run.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # linhas
    for ri, row_data in enumerate(rows):
        fill = "F5F5F5" if ri % 2 == 0 else "FFFFFF"
        tr   = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            c   = tr.cells[ci]
            shading(c, fill)
            cell_border(c)
            run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

def highlight_box(text, bg="EBF5EB", border_color="2E7D32"):
    table = doc.add_table(rows=1, cols=1)
    c     = table.rows[0].cells[0]
    shading(c, bg)
    cell_border(c)
    p     = c.paragraphs[0]
    run   = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ══════════════════════════════════════════════════════════════════════════════
# CAPA
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run("SISTEMA DE CONTROLE DE ESTOQUE")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = VERDE_ESCURO
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
run = p.add_run("Documentação Técnica e Gerencial")
run.font.size = Pt(16)
run.font.color.rgb = OURO
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
run = p.add_run("Versão 3.1.0")
run.font.size = Pt(13)
run.font.color.rgb = CINZA_ESCURO
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph().paragraph_format.space_after = Pt(8)
divider()
doc.add_paragraph().paragraph_format.space_after = Pt(8)

p = doc.add_paragraph()
run = p.add_run(f"Data de emissão: {datetime.date.today().strftime('%d de %B de %Y')}")
run.font.size = Pt(10)
run.font.color.rgb = CINZA_ESCURO
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
run = p.add_run("Responsável técnico: github.com/Wbad-02")
run.font.size = Pt(10)
run.font.color.rgb = CINZA_ESCURO
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 1. VISÃO GERAL
# ══════════════════════════════════════════════════════════════════════════════
heading("1. VISÃO GERAL DO SISTEMA", 1)
body(
    "O Sistema de Controle de Estoque é uma aplicação web full-stack desenvolvida para "
    "gerenciar materiais, patrimônio, ativos e movimentações de estoque em ambiente de "
    "rede interna corporativa. O sistema oferece rastreamento individual de itens por "
    "código patrimonial, atribuição de materiais a colaboradores/equipamentos, "
    "geração de relatórios em Excel e PDF, notificações automáticas por e-mail e "
    "importação de Notas Fiscais Eletrônicas (NF-e)."
)

heading("1.1 Informações do Produto", 2)
make_table(
    ["Atributo", "Valor"],
    [
        ["Nome",          "Sistema de Controle de Estoque"],
        ["Versão",        "3.1.0"],
        ["Tipo",          "Aplicação Web (SPA + REST API)"],
        ["Repositório",   "github.com/Wbad-02"],
        ["Ambiente",      "Rede interna corporativa (sem acesso externo)"],
        ["Banco de Dados","SQLite (arquivo local estoque.db)"],
        ["Servidor",      "FastAPI + Uvicorn (Python)"],
        ["Frontend",      "HTML5 + JavaScript Vanilla (SPA)"],
    ],
    [5, 11]
)

heading("1.2 Propósito e Escopo", 2)
bullet("Controle de entrada e saída de materiais de estoque")
bullet("Rastreamento patrimonial de unidades individuais (ex: monitores, notebooks)")
bullet("Atribuição de materiais a colaboradores ou equipamentos")
bullet("Alertas automáticos quando estoque atinge nível mínimo")
bullet("Exportação de relatórios em Excel e PDF")
bullet("Importação de itens via XML de Nota Fiscal Eletrônica (NF-e 4.0)")
bullet("Notificações por e-mail configuráveis")
bullet("Auditoria completa de todas as ações do sistema")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 2. TECNOLOGIAS
# ══════════════════════════════════════════════════════════════════════════════
heading("2. TECNOLOGIAS UTILIZADAS", 1)

heading("2.1 Backend", 2)
make_table(
    ["Biblioteca", "Versão", "Finalidade"],
    [
        ["FastAPI",              "0.115.14", "Framework web assíncrono para construção da REST API"],
        ["Uvicorn",              "0.34.3",   "Servidor ASGI de alta performance"],
        ["SQLAlchemy",           "2.0.49",   "ORM (Object-Relational Mapper) para acesso ao banco"],
        ["SQLite",               "built-in", "Banco de dados relacional embarcado"],
        ["python-jose",          "3.3.0",    "Geração e validação de tokens JWT"],
        ["passlib + bcrypt",     "1.7.4 / 4.0.1", "Hashing seguro de senhas (bcrypt, custo 12)"],
        ["OpenPyXL",             "3.1.2",    "Geração de planilhas Excel (.xlsx)"],
        ["ReportLab",            "4.1.0",    "Geração de documentos PDF"],
        ["python-multipart",     "0.0.9",    "Upload de arquivos (XML de NF-e)"],
    ],
    [4, 2.5, 9.5]
)

heading("2.2 Frontend", 2)
make_table(
    ["Tecnologia", "Uso"],
    [
        ["HTML5",             "Estrutura da interface (index.html — 973 linhas)"],
        ["JavaScript Vanilla","Toda lógica de cliente, SPA, chamadas de API (app.js — 2274 linhas)"],
        ["CSS3",              "Estilos, design system, responsividade (style.css — 403 linhas)"],
        ["LocalStorage",      "Persistência do token JWT e dados de sessão no navegador"],
        ["Fetch API",         "Comunicação assíncrona com a REST API"],
    ],
    [4.5, 11.5]
)

heading("2.3 Infraestrutura", 2)
make_table(
    ["Componente",  "Descrição"],
    [
        ["Servidor",    "Windows Server 2019+ ou Windows 10+ com Python 3.10+"],
        ["Porta",       "8000 (padrão Uvicorn, configurável)"],
        ["Rede",        "Interna corporativa — acesso bloqueado externamente por middleware e firewall"],
        ["Backup",      "Script PowerShell automatizado para cópia do estoque.db"],
        ["Deploy",      "Execução direta via Python (sem Docker, sem serviço cloud)"],
    ],
    [4, 12]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 3. ARQUITETURA
# ══════════════════════════════════════════════════════════════════════════════
heading("3. ARQUITETURA DO SISTEMA", 1)

heading("3.1 Visão Geral da Arquitetura", 2)
body(
    "O sistema segue uma arquitetura cliente-servidor em duas camadas: o frontend é "
    "um SPA (Single Page Application) servido como arquivos estáticos pelo próprio "
    "servidor FastAPI, e o backend expõe uma REST API que o frontend consome via HTTP."
)

for linha in [
    "Navegador (Chrome/Edge/Firefox)",
    "    ↕  HTTP/JSON",
    "FastAPI + Uvicorn  [porta 8000]",
    "    ├─ Middleware de Segurança (IP whitelist, rate limit, headers)",
    "    ├─ Routers REST API (/api/...)",
    "    ├─ SQLAlchemy ORM",
    "    └─ SQLite (estoque.db)",
]:
    code_block(linha)

heading("3.2 Estrutura de Diretórios", 2)
for linha in [
    "estoque_app/",
    "├── main.py                    # Ponto de entrada — monta app FastAPI",
    "├── models.py                  # Modelos SQLAlchemy (13 tabelas)",
    "├── schemas.py                 # Schemas Pydantic (validação entrada/saída)",
    "├── auth.py                    # JWT, bcrypt, decoradores de permissão",
    "├── database.py                # Engine SQLite + sessão SQLAlchemy",
    "├── middleware_seguranca.py    # IP whitelist, rate limiting, security headers",
    "├── email_service.py           # Envio de e-mails em background thread",
    "├── utils.py                   # sync_qty() — sincroniza quantidade de material",
    "├── requirements.txt           # Dependências Python",
    "├── estoque.db                 # Banco de dados SQLite",
    "├── smtp_config.json           # Config SMTP (gerado pela interface)",
    "├── routers/",
    "│   ├── usuarios.py            # CRUD usuários, alteração de senha",
    "│   ├── categorias.py          # Categorias de materiais",
    "│   ├── grupos.py              # Grupos dentro de categorias",
    "│   ├── materiais.py           # CRUD materiais + entradas",
    "│   ├── patrimonio.py          # Rastreamento de unidades individuais",
    "│   ├── retiradas.py           # Histórico de retiradas/saídas",
    "│   ├── ativos.py              # Ativos (colaboradores/equipamentos)",
    "│   ├── ativos_categorias.py   # Categorias e grupos de ativos",
    "│   ├── motivos.py             # Motivos personalizados de retirada",
    "│   ├── notificacoes.py        # Gerenciamento de e-mails e templates",
    "│   ├── relatorios.py          # Exportação Excel/PDF",
    "│   └── importacao.py          # Importação de NF-e (XML)",
    "├── static/",
    "│   ├── index.html             # Interface HTML (SPA)",
    "│   ├── app.js                 # Lógica JavaScript completa",
    "│   └── style.css              # Estilos e design system",
    "└── security/",
    "    ├── README_SEGURANCA.md    # Guia passo-a-passo de hardening",
    "    └── scripts PowerShell     # Configuração de rede, firewall e backup",
]:
    code_block(linha)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 4. BANCO DE DADOS
# ══════════════════════════════════════════════════════════════════════════════
heading("4. BANCO DE DADOS", 1)

heading("4.1 Configuração", 2)
make_table(
    ["Parâmetro", "Valor"],
    [
        ["Motor",           "SQLite"],
        ["ORM",             "SQLAlchemy 2.0"],
        ["Modo WAL",        "Ativado (melhor concorrência de leitura/escrita)"],
        ["Synchronous",     "NORMAL (equilíbrio entre segurança e performance)"],
        ["Thread-safe",     "check_same_thread=False (gerenciado pelo ORM)"],
        ["Arquivo",         "estoque.db (mesmo diretório da aplicação)"],
    ],
    [5, 11]
)

heading("4.2 Modelo de Dados — Tabelas", 2)
make_table(
    ["Tabela", "Descrição", "Campos Principais"],
    [
        ["usuarios",               "Contas de acesso ao sistema",
         "id, nome, email, senha_hash, grupo(enum), ativo, criado_em"],
        ["categorias",             "Categorias de materiais",
         "id, nome, descricao, criado_em"],
        ["grupos_material",        "Grupos dentro de categorias",
         "id, nome, descricao, quantidade_minima, categoria_id, criado_em"],
        ["materiais",              "Produtos em estoque",
         "id, nome, descricao, quantidade, unidade, grupo_id, ativo, usa_patrimonio, valor_unitario, tag"],
        ["unidades_patrimonio",    "Rastreamento de unidades individuais",
         "id, material_id, codigo, status(enum), origem, nf_numero, valor_unitario, tag, criado_em, retirado_em"],
        ["movimentacoes",          "Entradas e saídas de estoque",
         "id, material_id, usuario_id, unidade_id, tipo(entrada/saida), quantidade, motivo, observacao, criado_em"],
        ["ativos_categorias",      "Categorias de ativos (ex: Colaboradores)",
         "id, nome, descricao, criado_em"],
        ["ativos_grupos",          "Grupos dentro de categorias de ativos",
         "id, nome, descricao, categoria_id, criado_em"],
        ["ativos",                 "Entidades que recebem materiais",
         "id, nome, descricao, grupo_id, ativo, criado_em"],
        ["ativos_itens",           "Materiais atribuídos a ativos",
         "id, ativo_id, material_id, unidade_id, quantidade, observacao, atribuido_em, devolvido_em"],
        ["logs_auditoria",         "Trilha completa de ações do sistema",
         "id, usuario_id, acao, entidade, entidade_id, detalhe, criado_em"],
        ["motivos_personalizados", "Motivos customizados de retirada",
         "id, nome, ativo, criado_em"],
        ["notificacoes_emails",    "E-mails para notificações automáticas",
         "id, email, tipo, ativo, intervalo_dias, criado_em"],
        ["notificacoes_templates", "Templates de e-mail por tipo",
         "id, tipo, assunto, corpo, atualizado_em"],
    ],
    [3.5, 4.5, 8]
)

heading("4.3 Hierarquia de Dados", 2)
for linha in [
    "Categoria  (ex: 'Informática')",
    "  └─ Grupo  (ex: 'Monitores' — define quantidade_minima)",
    "      └─ Material  (ex: 'Monitor Dell 24' — quantidade float)",
    "         └─ [se usa_patrimonio=true]",
    "            └─ UnidadePatrimônio  (ex: TI-001, TI-002...)",
    "               — rastreia: código, origem, NF, status, retirada",
]:
    code_block(linha)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 5. AUTENTICAÇÃO E CONTROLE DE ACESSO
# ══════════════════════════════════════════════════════════════════════════════
heading("5. AUTENTICAÇÃO E CONTROLE DE ACESSO", 1)

heading("5.1 Mecanismo de Autenticação", 2)
make_table(
    ["Componente", "Detalhe"],
    [
        ["Algoritmo de senha", "bcrypt com fator de custo 12 (resistente a brute-force)"],
        ["Token de sessão",    "JWT (JSON Web Token) — algoritmo HS256"],
        ["Validade do token",  "8 horas (expiração automática)"],
        ["Chave secreta",      "Variável de ambiente ESTOQUE_SECRET_KEY (fallback: gerada aleatória por restart)"],
        ["Renovação",          "Requer novo login após expiração"],
        ["Armazenamento",      "Token em localStorage do navegador"],
        ["Auto-logout",        "Frontend detecta resposta 401 e desloga automaticamente"],
    ],
    [5, 11]
)

heading("5.2 Hierarquia de Grupos de Acesso", 2)
make_table(
    ["Grupo", "Nível", "Descrição", "Permissões Especiais"],
    [
        ["mestre", "4 (máx.)",  "Acesso total — oculto na interface",
         "Criar/editar admins, redefinir qualquer senha"],
        ["admin",  "3",         "Administrador do sistema",
         "Criar usuários (editor/viewer), gerenciar notificações, importar NF-e, excluir materiais"],
        ["editor", "2",         "Editor de conteúdo",
         "Criar/editar materiais, registrar retiradas, gerenciar ativos e categorias"],
        ["viewer", "1 (mín.)",  "Somente leitura",
         "Dashboard, listagens, histórico, exportar relatórios"],
    ],
    [2.5, 2.5, 5, 6]
)

heading("5.3 Matriz de Permissões por Funcionalidade", 2)
make_table(
    ["Funcionalidade", "viewer", "editor", "admin", "mestre"],
    [
        ["Ver dashboard e materiais",       "✓", "✓", "✓", "✓"],
        ["Ver histórico de retiradas",      "✓", "✓", "✓", "✓"],
        ["Exportar relatórios (Excel/PDF)", "✓", "✓", "✓", "✓"],
        ["Registrar entrada de material",   "—", "✓", "✓", "✓"],
        ["Registrar saída/retirada",        "—", "✓", "✓", "✓"],
        ["Criar/editar material",           "—", "✓", "✓", "✓"],
        ["Excluir material",                "—", "—", "✓", "✓"],
        ["Gerenciar categorias/grupos",     "—", "✓", "✓", "✓"],
        ["Gerenciar ativos",                "—", "✓", "✓", "✓"],
        ["Criar/editar usuários",           "—", "—", "✓", "✓"],
        ["Gerenciar notificações",          "—", "—", "✓", "✓"],
        ["Exportar relatório notificações", "—", "—", "✓", "✓"],
        ["Importar NF-e (XML)",             "—", "—", "✓", "✓"],
        ["Criar admin / resetar senhas",    "—", "—", "—", "✓"],
    ],
    [7, 1.8, 1.8, 1.8, 1.8]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 6. SEGURANÇA
# ══════════════════════════════════════════════════════════════════════════════
heading("6. SEGURANÇA", 1)

heading("6.1 Camadas de Proteção", 2)
body("O sistema implementa 5 camadas de proteção configuradas no middleware_seguranca.py:")

make_table(
    ["Camada", "Mecanismo", "Proteção Oferecida"],
    [
        ["1 — Whitelist de IP", "CIDR allowlist por IP/sub-rede",
         "Bloqueia (403) qualquer acesso de fora da rede interna"],
        ["2 — Rate Limit Geral", "60 req/60s por IP (janela deslizante)",
         "Previne sobrecarga e scraping automático (429)"],
        ["3 — Rate Limit Login", "10 tentativas/5 min por IP",
         "Bloqueia ataque de força bruta nas credenciais (429)"],
        ["4 — Security Headers", "7 headers HTTP de segurança",
         "Protege contra XSS, clickjacking, MIME sniffing e cache indevido"],
        ["5 — Bloqueio de Métodos", "Allowlist de verbos HTTP",
         "Nega TRACE, CONNECT e outros métodos não necessários"],
    ],
    [4, 5, 7]
)

heading("6.2 Security Headers Configurados", 2)
make_table(
    ["Header", "Valor", "Finalidade"],
    [
        ["X-Content-Type-Options",  "nosniff",       "Impede MIME type sniffing"],
        ["X-Frame-Options",         "DENY",           "Bloqueia clickjacking via iframe"],
        ["X-XSS-Protection",        "1; mode=block",  "Ativa proteção XSS do navegador"],
        ["Referrer-Policy",         "no-referrer",    "Não vaza URL de origem em requests"],
        ["Content-Security-Policy", "self + unsafe-inline", "Restringe fontes de conteúdo"],
        ["Permissions-Policy",      "geolocation=(), camera=(), microphone=()", "Desativa APIs sensíveis do browser"],
        ["Cache-Control",           "no-store, no-cache", "Impede cache de dados sensíveis"],
        ["Server",                  "Estoque/3.0",    "Oculta tecnologia real do servidor"],
    ],
    [5, 4.5, 6.5]
)

highlight_box(
    "IMPORTANTE — Credencial Padrão: Na primeira inicialização, o sistema cria automaticamente "
    "o usuário admin@estoque.local com a senha admin123. Esta senha DEVE ser alterada "
    "imediatamente após o primeiro acesso para garantir a segurança do sistema."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 7. FUNCIONALIDADES
# ══════════════════════════════════════════════════════════════════════════════
heading("7. FUNCIONALIDADES DO SISTEMA", 1)

heading("7.1 Gerenciamento de Estoque", 2)
bullet("Cadastro de materiais com categorias e grupos hierárquicos")
bullet("Controle de quantidade com alertas de estoque mínimo por grupo")
bullet("Registro de entradas com quantidade, observação e valor unitário")
bullet("Registro de saídas/retiradas com motivo, observação e usuário responsável")
bullet("Motivos de retirada padrão (colaborador, defeito) + customizáveis")
bullet("Timeline completa de movimentações por material")
bullet("Dashboard com contadores em tempo real e filtros por categoria/grupo")

heading("7.2 Rastreamento Patrimonial", 2)
bullet("Materiais podem ter rastreamento individual por unidade (usa_patrimonio=true)")
bullet("Cada unidade recebe código patrimonial único (ex: TI-001, MON-005)")
bullet("Rastreamento de origem: manual, importação XML, ou sistema")
bullet("Histórico completo: data de entrada, retirada, quem retirou, motivo")
bullet("Sincronização automática: quantidade do material = total de unidades ativas")
bullet("Importação em lote via NF-e: cria unidades com número de NF e valor unitário")

heading("7.3 Gestão de Ativos", 2)
bullet("Ativos representam colaboradores, departamentos ou equipamentos")
bullet("Materiais podem ser atribuídos a ativos (com quantidade e observação)")
bullet("Unidades atribuídas são excluídas da contagem de disponíveis automaticamente")
bullet("Devolução registra data e restaura material ao estoque")
bullet("Inativação de ativo retorna todos os materiais ao estoque automaticamente")
bullet("Histórico de todos os materiais já associados a cada ativo")

heading("7.4 Relatórios e Exportação", 2)
make_table(
    ["Relatório", "Formato", "Conteúdo", "Acesso"],
    [
        ["Estoque Completo",  "Excel / PDF", "Todos materiais ativos: categoria, grupo, qtd, unidade, status", "Todos"],
        ["Apenas Alertas",    "Excel / PDF", "Materiais abaixo ou igual ao mínimo do grupo",                   "Todos"],
        ["Saídas",            "Excel",       "Histórico de retiradas com filtros: motivo, período",             "Todos"],
        ["Ativos",            "Excel / PDF", "Lista de ativos com status: ativo, inativo ou todos",             "Todos"],
        ["Notificações",      "Excel",       "E-mails cadastrados para notificações por tipo",                  "Admin+"],
    ],
    [3.5, 2.5, 7.5, 2.5]
)

heading("7.5 Notificações por E-mail", 2)
make_table(
    ["Tipo",      "Disparado Quando",                  "Configurável"],
    [
        ["retirada", "Saída de material é registrada",        "Sim (e-mails, template)"],
        ["entrada",  "Entrada de material é registrada",       "Sim (e-mails, template)"],
        ["alerta",   "Quantidade ≤ mínimo do grupo",           "Sim (e-mails, template, intervalo em dias)"],
    ],
    [3, 7.5, 5.5]
)
body("Envio em background thread (não bloqueia a API). Templates com variáveis: {material}, {quantidade}, {usuario}, {data}, etc.")

heading("7.6 Importação de NF-e", 2)
bullet("Upload de arquivo XML de NF-e versão 4.0")
bullet("Preview automático: emitente, número da NF, itens com código, nome, qtd e valor")
bullet("Importação cria materiais novos (se não existirem) e suas unidades patrimoniais")
bullet("Unidades importadas ficam marcadas com origem=xml e nf_numero da nota")
bullet("Disponível apenas para admin e mestre")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 8. INTERFACE DO USUÁRIO
# ══════════════════════════════════════════════════════════════════════════════
heading("8. INTERFACE DO USUÁRIO", 1)

heading("8.1 Estrutura de Navegação", 2)
make_table(
    ["Seção", "Visível para", "Descrição"],
    [
        ["Dashboard",        "Todos",      "Visão geral: totais, alertas, tabela de materiais com filtros"],
        ["Materiais",        "Editor+",    "CRUD de materiais, timeline, entrada/saída, rastreamento patrimonial"],
        ["Ativos",           "Editor+",    "Gestão de ativos, atribuição e devolução de materiais"],
        ["Cadastros",        "Editor+",    "Categorias e grupos de materiais e ativos"],
        ["Retiradas",        "Editor+",    "Formulário de retirada + histórico completo com filtros"],
        ["Notificações",     "Admin+",     "Gerenciamento de e-mails e templates de notificação"],
        ["Usuários",         "Admin+",     "CRUD de usuários, reset de senhas"],
        ["Importar NF-e",    "Admin+",     "Upload e importação de XML de NF-e"],
        ["Relatórios",       "Todos",      "Geração de relatórios Excel/PDF por seção"],
        ["Perfil",           "Todos",      "Edição de nome e troca de senha do usuário logado"],
    ],
    [3.5, 2.5, 10]
)

heading("8.2 Design System", 2)
make_table(
    ["Elemento", "Especificação"],
    [
        ["Paleta principal",  "Verde escuro #1E3A34, Ouro #B8952A, Cinza claro #F5F5F5"],
        ["Tipografia",        "System-UI sans-serif, escala de 11px a 20px"],
        ["Responsividade",    "Breakpoint em 1024px — sidebar recolhível em telas menores"],
        ["Feedback visual",   "Toast notifications, badges de status, loading states"],
        ["Componentes",       "Modais, tabelas com zebra striping, cards, dropdowns, badges"],
    ],
    [4, 12]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 9. ENDPOINTS DA API
# ══════════════════════════════════════════════════════════════════════════════
heading("9. ENDPOINTS DA API REST", 1)
body("Todos os endpoints são prefixados com /api. Autenticação via header Authorization: Bearer {token}.")

heading("9.1 Autenticação", 2)
make_table(
    ["Método", "Endpoint",          "Descrição",                       "Requer"],
    [
        ["POST", "/auth/login",     "Login — retorna JWT + grupo",     "Público"],
    ],
    [1.5, 4, 8.5, 2]
)

heading("9.2 Materiais", 2)
make_table(
    ["Método", "Endpoint",                         "Descrição",                          "Requer"],
    [
        ["GET",    "/materiais/",                  "Listar materiais (filtros: grupo, categoria, alertas)", "Autenticado"],
        ["POST",   "/materiais/",                  "Criar material",                     "Editor+"],
        ["GET",    "/materiais/{id}",              "Detalhe de material",                "Autenticado"],
        ["PUT",    "/materiais/{id}",              "Editar material",                    "Editor+"],
        ["DELETE", "/materiais/{id}",              "Remover material (soft delete)",     "Admin+"],
        ["POST",   "/materiais/{id}/entrada",      "Registrar entrada de estoque",       "Editor+"],
        ["GET",    "/materiais/alertas/count",     "Contar materiais em alerta",         "Autenticado"],
    ],
    [1.5, 5, 6.5, 3]
)

heading("9.3 Patrimônio, Retiradas e Ativos", 2)
make_table(
    ["Método", "Endpoint",                                    "Descrição",                          "Requer"],
    [
        ["GET",    "/patrimonio/{mat_id}/unidades",           "Listar unidades de um material",     "Autenticado"],
        ["POST",   "/patrimonio/{mat_id}/unidades",           "Adicionar unidades patrimoniais",    "Editor+"],
        ["PATCH",  "/patrimonio/{mat_id}/unidades/{uid}/codigo","Editar código patrimonial",        "Editor+"],
        ["GET",    "/retiradas/",                             "Histórico de retiradas (paginado)",  "Autenticado"],
        ["GET",    "/ativos/",                                "Listar ativos ativos",               "Autenticado"],
        ["POST",   "/ativos/",                                "Criar ativo",                        "Editor+"],
        ["POST",   "/ativos/{id}/atribuir",                   "Atribuir material a ativo",          "Editor+"],
        ["POST",   "/ativos/{id}/devolver/{item_id}",         "Devolver material de ativo",         "Editor+"],
        ["DELETE", "/ativos/{id}",                            "Inativar ativo (retorna materiais)", "Editor+"],
        ["POST",   "/ativos/{id}/reativar",                   "Reativar ativo",                     "Editor+"],
    ],
    [1.5, 6, 5.5, 3]
)

heading("9.4 Relatórios, Notificações e Importação", 2)
make_table(
    ["Método", "Endpoint",                           "Descrição",                           "Requer"],
    [
        ["GET", "/relatorios/excel",                 "Exportar estoque em Excel",           "Autenticado"],
        ["GET", "/relatorios/pdf",                   "Exportar estoque em PDF",             "Autenticado"],
        ["GET", "/relatorios/saidas/excel",          "Exportar retiradas em Excel",         "Autenticado"],
        ["GET", "/relatorios/ativos/excel",          "Exportar ativos em Excel/PDF",        "Autenticado"],
        ["GET", "/relatorios/notificacoes/excel",    "Exportar e-mails de notificação",     "Admin+"],
        ["GET", "/notificacoes/emails",              "Listar e-mails cadastrados",          "Admin+"],
        ["POST","/notificacoes/emails",              "Cadastrar e-mail de notificação",     "Admin+"],
        ["PUT", "/notificacoes/templates/{tipo}",    "Atualizar template de e-mail",        "Admin+"],
        ["GET", "/notificacoes/smtp",                "Obter config SMTP",                   "Admin+"],
        ["POST","/notificacoes/smtp",                "Salvar config SMTP",                  "Admin+"],
        ["POST","/importacao/preview",               "Preview de XML NF-e",                 "Admin+"],
        ["POST","/importacao/importar",              "Importar itens do XML NF-e",          "Admin+"],
    ],
    [1.5, 5.5, 6, 3]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 10. FLUXOS PRINCIPAIS
# ══════════════════════════════════════════════════════════════════════════════
heading("10. FLUXOS PRINCIPAIS DE USO", 1)

heading("10.1 Login e Sessão", 2)
for i, passo in enumerate([
    "Usuário acessa http://<IP_DO_SERVIDOR>:8000 pelo navegador",
    "Se não há token válido em localStorage → tela de login exibida",
    "Usuário informa e-mail e senha → POST /api/auth/login",
    "API valida credenciais (bcrypt), gera JWT com validade de 8 horas",
    "Frontend armazena token, grupo e nome no localStorage",
    "Todas as requisições incluem o header Authorization: Bearer {token}",
    "Se a API retornar 401 (token expirado/inválido) → logout automático",
], 1):
    bullet(f"{i}. {passo}")

heading("10.2 Registrar Entrada de Material", 2)
for i, passo in enumerate([
    "Editor acessa a seção Materiais",
    "Clica em '+ Adicionar' ou em um material existente → opção de Entrada",
    "Informa quantidade, observação e valor unitário (opcional)",
    "Sistema registra Movimentacao (tipo=entrada) e atualiza o campo quantidade",
    "Se notificações de entrada configuradas → e-mail disparado em background",
    "Dashboard e listagem atualizam a quantidade em tempo real",
], 1):
    bullet(f"{i}. {passo}")

heading("10.3 Registrar Retirada (Saída)", 2)
for i, passo in enumerate([
    "Editor acessa Retiradas e preenche: categoria, grupo, material, motivo e observação",
    "Se material usa patrimônio: seleciona unidade específica a ser retirada",
    "Confirma → API registra Movimentacao (tipo=saida)",
    "Se unidade patrimonial: status=retirado, registra data e usuário",
    "sync_qty() recalcula quantidade do material automaticamente",
    "Se quantidade caiu abaixo do mínimo → alerta no dashboard e e-mail de alerta",
    "Log de auditoria registra quem retirou, quando, quanto e por quê",
], 1):
    bullet(f"{i}. {passo}")

heading("10.4 Importar NF-e", 2)
for i, passo in enumerate([
    "Admin acessa 'Importar NF-e' e faz upload do arquivo XML",
    "API analisa XML NF-e 4.0, extrai emitente, número da NF e itens",
    "Tela exibe preview com todos os itens encontrados",
    "Admin confirma a importação",
    "Para cada item: cria Material (se não existir) e UnidadePatrimônio (origem=xml)",
    "Unidades ficam marcadas com número da NF e valor unitário da nota",
    "Quantidade do material é sincronizada automaticamente",
], 1):
    bullet(f"{i}. {passo}")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 11. PONTOS POSITIVOS E NEGATIVOS
# ══════════════════════════════════════════════════════════════════════════════
heading("11. ANÁLISE CRÍTICA: PONTOS POSITIVOS E NEGATIVOS", 1)

heading("11.1 Pontos Positivos", 2, color=VERDE_MEDIO)
make_table(
    ["Ponto Positivo", "Detalhe"],
    [
        ["Segurança em camadas",
         "5 camadas independentes: IP whitelist, rate limiting duplo, security headers, "
         "bcrypt custo 12 e JWT com expiração. Proteção sólida para ambiente interno."],
        ["Zero dependências de infraestrutura",
         "SQLite embarcado, sem servidor de BD, sem Redis, sem serviços externos. "
         "Deploy em qualquer Windows com Python instalado em minutos."],
        ["Rastreamento patrimonial completo",
         "Cada unidade tem código, origem, NF, histórico de retirada e sincronização "
         "automática de quantidade. Muito além de um simples controle de quantidade."],
        ["Auditoria total",
         "Toda ação sensível é registrada com usuário, timestamp e detalhe. "
         "Permite rastrear qualquer alteração histórica."],
        ["Controle de acesso granular",
         "4 níveis bem definidos (viewer/editor/admin/mestre) com hierarquia clara "
         "e aplicada tanto no backend (decoradores) quanto no frontend (CSS classes)."],
        ["Importação de NF-e",
         "Elimina trabalho manual: importa itens diretamente do XML da nota fiscal, "
         "criando materiais e unidades patrimoniais automaticamente."],
        ["Notificações configuráveis",
         "Alertas automáticos por e-mail para retiradas, entradas e estoque mínimo, "
         "com templates editáveis e intervalos configuráveis."],
        ["Relatórios em Excel e PDF",
         "Exportação com formatação profissional, filtros, cores de alerta e timestamp "
         "no nome do arquivo para controle de versão."],
        ["Frontend sem framework",
         "Sem dependências de npm, React ou Angular. Zero vulnerabilidades de "
         "supply chain via pacotes JavaScript. Simples de manter."],
        ["Documentação de segurança",
         "Guia completo em security/README_SEGURANCA.md com passos práticos "
         "para hardening, firewall, IP estático e backup."],
    ],
    [5, 11]
)

heading("11.2 Pontos Negativos / Riscos", 2, color=VERMELHO)
make_table(
    ["Ponto Negativo / Risco", "Impacto", "Recomendação"],
    [
        ["Rate limiting em memória",
         "Médio — reinicia com o servidor, não persiste entre reinicializações.",
         "Migrar para Redis ou persistir contadores em banco para ambientes de alta criticidade."],
        ["SQLite para múltiplos usuários",
         "Médio — adequado até ~100 usuários simultâneos. Pode apresentar locks em escrita sob carga alta.",
         "Para mais de 50 usuários ativos: migrar para PostgreSQL (SQLAlchemy facilita a troca)."],
        ["Sem autenticação de dois fatores (2FA)",
         "Médio — acesso apenas por senha. Se credenciais vazarem, a conta fica exposta.",
         "Implementar TOTP (Google Authenticator) ou envio de código por e-mail."],
        ["Logs de auditoria sem interface",
         "Baixo — tabela logs_auditoria existe no banco mas não há tela para consulta.",
         "Adicionar página de auditoria para admin com filtros por usuário, ação e período."],
        ["Sem backup automático nativo",
         "Alto — perda do arquivo estoque.db = perda total dos dados.",
         "Configurar imediatamente o script de backup do PowerShell (security/backup_automatico.ps1)."],
        ["Chave JWT sem persistência",
         "Médio — sem ESTOQUE_SECRET_KEY definida, cada reinício invalida todos os tokens.",
         "Definir ESTOQUE_SECRET_KEY como variável de ambiente permanente no servidor."],
        ["Ausência de testes automatizados na codebase principal",
         "Médio — alterações podem introduzir regressões sem ser detectadas rapidamente.",
         "Manter e expandir o _test_suite.py para cobertura de todos os fluxos críticos."],
        ["SMTP sem TLS obrigatório",
         "Baixo — configuração SMTP pode ser criada sem TLS ativado.",
         "Validar que smtp_config.json sempre use TLS=true em produção."],
        ["Frontend monolítico",
         "Baixo — app.js com 2274 linhas dificulta manutenção conforme o sistema cresce.",
         "Modularizar em arquivos separados por funcionalidade conforme o sistema evoluir."],
        ["Sem suporte a múltiplos depósitos/filiais",
         "Baixo — sistema assume depósito único.",
         "Para empresas com múltiplas filiais, seria necessário redesenho do modelo de dados."],
    ],
    [5, 3.5, 7.5]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 12. OPERAÇÃO E MANUTENÇÃO
# ══════════════════════════════════════════════════════════════════════════════
heading("12. OPERAÇÃO E MANUTENÇÃO", 1)

heading("12.1 Instalação", 2)
for i, passo in enumerate([
    "Instalar Python 3.10+ no servidor Windows",
    "Copiar a pasta estoque_app para o servidor",
    "Executar: pip install -r requirements.txt",
    "Definir variável de ambiente: set ESTOQUE_SECRET_KEY=<chave_aleatória>",
    "Iniciar servidor: python main.py  (ou uvicorn main:app --host 0.0.0.0 --port 8000)",
    "Acessar no navegador: http://<IP_DO_SERVIDOR>:8000",
    "Fazer login com admin@estoque.local / admin123 e alterar a senha",
], 1):
    bullet(f"{i}. {passo}")

heading("12.2 Backup do Banco de Dados", 2)
body("O banco de dados é um único arquivo SQLite (estoque.db). Para fazer backup:")
bullet("Manual: copiar o arquivo estoque.db para local seguro")
bullet("Automático: usar o script backup_automatico.ps1 em security/ com o Agendador de Tarefas do Windows")
bullet("Retenção recomendada: mínimo 30 dias de histórico de backups")

highlight_box(
    "CRÍTICO — Configure o backup automático antes de colocar o sistema em produção. "
    "A perda do arquivo estoque.db resulta em perda total de todos os dados do sistema."
)

heading("12.3 Monitoramento", 2)
make_table(
    ["O que monitorar", "Como verificar", "Ação se detectado"],
    [
        ["Tentativas de login bloqueadas", "Logs stdout da aplicação (429)",
         "Investigar IP de origem — possível ataque de força bruta"],
        ["Acessos bloqueados (403)",       "Logs stdout da aplicação",
         "Verificar se é IP legítimo não cadastrado na whitelist"],
        ["Alertas de estoque mínimo",      "Dashboard — badge ⚠ / e-mail",
         "Registrar entrada de material para o item em alerta"],
        ["Erros de envio de e-mail",       "Logs stdout da aplicação",
         "Verificar configuração SMTP em Notificações"],
        ["Espaço em disco do servidor",    "Monitoramento do SO",
         "Arquivos de backup e o estoque.db crescem com o tempo"],
    ],
    [5, 5, 6]
)

heading("12.4 Tarefas Periódicas Recomendadas", 2)
make_table(
    ["Frequência",  "Tarefa"],
    [
        ["Diário",     "Verificar alertas de estoque mínimo no dashboard"],
        ["Semanal",    "Revisar logs de auditoria para ações incomuns"],
        ["Mensal",     "Revisar usuários inativos e revogar acessos desnecessários"],
        ["Mensal",     "Testar restauração de backup (copiar estoque.db para ambiente de teste)"],
        ["Trimestral", "Atualizar dependências Python (pip install -r requirements.txt --upgrade)"],
        ["Trimestral", "Revisar whitelist de IP se a rede interna mudou"],
    ],
    [3, 13]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 13. CHECKLIST PÓS-DEPLOY
# ══════════════════════════════════════════════════════════════════════════════
heading("13. CHECKLIST PÓS-DEPLOY", 1)

heading("13.1 Segurança Obrigatória", 2)
make_table(
    ["#", "Item",                                                    "Status"],
    [
        ["1", "Senha do admin padrão alterada (admin@estoque.local)",  "[ ]"],
        ["2", "ESTOQUE_SECRET_KEY definida como variável de ambiente", "[ ]"],
        ["3", "Whitelist de IP configurada com sub-rede correta",      "[ ]"],
        ["4", "Firewall do Windows bloqueando porta 8000 de fora",     "[ ]"],
        ["5", "SMTP configurado com TLS habilitado",                   "[ ]"],
    ],
    [0.8, 13.2, 1.5]
)

heading("13.2 Operação Básica", 2)
make_table(
    ["#", "Item",                                                             "Status"],
    [
        ["6",  "Backup automático agendado (Agendador de Tarefas)",            "[ ]"],
        ["7",  "Categorias e grupos de materiais criados",                     "[ ]"],
        ["8",  "E-mails de notificação cadastrados (admin)",                   "[ ]"],
        ["9",  "Usuários de acesso criados com grupos adequados",              "[ ]"],
        ["10", "Testado acesso de múltiplos computadores da rede",             "[ ]"],
        ["11", "Testado bloqueio de acesso externo (via dados móveis/hotspot)","[ ]"],
        ["12", "IP do servidor configurado como estático",                     "[ ]"],
    ],
    [0.8, 13.2, 1.5]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 14. MELHORIAS SUGERIDAS
# ══════════════════════════════════════════════════════════════════════════════
heading("14. ROADMAP — MELHORIAS SUGERIDAS", 1)

heading("14.1 Curto Prazo (Alta Prioridade)", 2, color=VERMELHO)
make_table(
    ["Melhoria",                          "Benefício",                                        "Complexidade"],
    [
        ["Tela de auditoria para admin",  "Rastreabilidade completa via interface sem acesso ao banco", "Baixa"],
        ["Persistência de rate limiting", "Proteção contra brute-force sobrevive a reinicializações",  "Média"],
        ["Exportação de auditoria CSV",   "Permite análise externa em Excel/BI",                       "Baixa"],
        ["Backup automático via interface","Admin agenda e monitora backups sem PowerShell",             "Média"],
    ],
    [5.5, 7.5, 3]
)

heading("14.2 Médio Prazo", 2, color=OURO)
make_table(
    ["Melhoria",                          "Benefício",                                        "Complexidade"],
    [
        ["Autenticação 2FA (TOTP)",       "Segurança adicional contra roubo de credenciais",          "Média"],
        ["Suporte a PostgreSQL",          "Escala para centenas de usuários simultâneos",              "Média"],
        ["Dashboard com gráficos",        "Visualização de tendências de consumo e movimentação",      "Média"],
        ["Modularização do frontend",     "Manutenção mais fácil conforme sistema cresce",             "Alta"],
        ["API pública documentada (Swagger habilitado)", "Integração com outros sistemas internos",    "Baixa"],
    ],
    [5.5, 7.5, 3]
)

heading("14.3 Longo Prazo", 2, color=CINZA_ESCURO)
make_table(
    ["Melhoria",                             "Benefício",                                     "Complexidade"],
    [
        ["Suporte a múltiplos depósitos",    "Empresas com filiais ou múltiplos almoxarifados","Alta"],
        ["App mobile (PWA ou React Native)", "Acesso via smartphone para inventário físico",   "Alta"],
        ["Integração com ERP/SAP",           "Sincronização automática com sistema principal", "Alta"],
        ["Fila de mensagens (Celery/RabbitMQ)","E-mails resilientes com retry automático",     "Alta"],
        ["Relatórios agendados automáticos", "Relatório semanal enviado por e-mail automaticamente","Média"],
    ],
    [5.5, 7.5, 3]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 15. GLOSSÁRIO
# ══════════════════════════════════════════════════════════════════════════════
heading("15. GLOSSÁRIO", 1)

make_table(
    ["Termo",                   "Definição"],
    [
        ["API REST",            "Interface de comunicação HTTP padronizada entre frontend e backend"],
        ["Ativo",               "Entidade (colaborador ou equipamento) que pode receber materiais atribuídos"],
        ["bcrypt",              "Algoritmo de hashing de senhas resistente a ataques de força bruta"],
        ["CIDR",                "Notação de faixa de endereços IP (ex: 192.168.1.0/24 = rede de 254 hosts)"],
        ["FastAPI",             "Framework Python moderno para criação de APIs REST assíncronas"],
        ["JWT",                 "JSON Web Token — credencial compacta e assinada para autenticação"],
        ["NF-e",                "Nota Fiscal eletrônica — documento fiscal em formato XML padrão SEFAZ"],
        ["ORM",                 "Mapeador Objeto-Relacional — camada que converte objetos Python em tabelas SQL"],
        ["Patrimônio",          "Unidade individual rastreável de um material (ex: monitor TI-001)"],
        ["Rate Limiting",       "Limite de requisições por tempo para proteger contra uso abusivo"],
        ["SPA",                 "Single Page Application — frontend que troca apenas partes da página sem reload"],
        ["soft delete",         "Exclusão lógica (marca ativo=False) sem remover o registro do banco"],
        ["SQLite",              "Banco de dados relacional embarcado em arquivo único, sem servidor"],
        ["sync_qty()",          "Função interna que recalcula a quantidade de um material baseado em suas unidades patrimoniais ativas"],
        ["WAL",                 "Write-Ahead Logging — modo SQLite que permite leituras simultâneas durante escrita"],
        ["Whitelist",           "Lista de IPs autorizados — qualquer IP fora da lista é bloqueado"],
    ],
    [3.5, 12.5]
)


# ══════════════════════════════════════════════════════════════════════════════
# RODAPÉ / ASSINATURA
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
divider()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
run = p.add_run("Sistema de Controle de Estoque v3.1.0")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = VERDE_ESCURO

p = doc.add_paragraph()
run = p.add_run(f"Documentação gerada em {datetime.date.today().strftime('%d/%m/%Y')}  |  github.com/Wbad-02")
run.font.size = Pt(9)
run.font.color.rgb = CINZA_ESCURO

p = doc.add_paragraph()
run = p.add_run("© Todos os direitos reservados – uso interno")
run.font.size = Pt(9)
run.font.color.rgb = CINZA_ESCURO
run.italic = True


# ── Salvar ──────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Documento salvo em: {OUTPUT}")
